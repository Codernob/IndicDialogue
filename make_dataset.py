from tqdm import tqdm
import pandas as pd
from pysubparser import parser
import os
import shutil
from pysubparser.cleaners import  brackets, formatting,lower_case,urls_cleaner
import argparse
from docx import Document


language_full_names = {
    "asm": "Assamese",
    "ben": "Bengali",
    "hin": "Hindi",
    "mar": "Marathi",
    "npl": "Nepali",
    "ori": "Odia",
    "snd": "Sindhi",
    "tam": "Tamil",
    "tel": "Telugu",
    "urd": "Urdu"
}
stats={
    'dialogs':0,
    'total_words':0,
    'total_characters':0,
}
stats_language_wise={
    k:{
        'dialogs':0,
        'total_words':0,
        'total_characters':0
    } for k in language_full_names.keys()
}
def create_report(total_stats:dict,language_stats:dict):
    # Create a new Document
    doc = Document()

    # Add language statistics
    doc.add_heading('Language Statistics', level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Language'
    hdr_cells[1].text = 'Dialogs'
    hdr_cells[2].text = 'Total Words'
    hdr_cells[3].text = 'Total Characters'
    hdr_cells[4].text = 'Files'
    for lang, stats in language_stats.items():
        row_cells = table.add_row().cells
        row_cells[0].text = lang
        row_cells[1].text = str(stats['dialogs'])
        row_cells[2].text = str(stats['total_words'])
        row_cells[3].text = str(stats['total_characters'])
        row_cells[4].text = str(stats['files'])

    # Add total statistics as the last row in the table
    total_row = table.add_row().cells
    total_row[0].text = 'Total'
    total_row[1].text = str(total_stats['dialogs'])
    total_row[2].text = str(total_stats['total_words'])
    total_row[3].text = str(total_stats['total_characters'])
    total_row[4].text = str(total_stats['files'])

    # Save the document
    doc.save('language_statistics.docx')


def find_file_path(file_path:list[str],id:str)->str:
    file_path="files//"+"//".join(file_path)
    final_file_path=file_path+"//"+id+".srt"
    return final_file_path
def delete_directory(directory_path):
    if os.path.exists(directory_path):
        try:
            print(f"Found existing directory.Removing...")
            shutil.rmtree(directory_path)
            print(f"Directory '{directory_path}' and its contents have been successfully deleted.")
        except OSError as e:
            print(f"Error: {directory_path} : {e.strerror}")
    else:
        print(f"Directory '{directory_path}' does not exist.Creating new one...")



def reverse_id(id:str)->list[str]:
    return [i for i  in id[:-5:-1]]

if __name__=='__main__':
    
    argument_parser = argparse.ArgumentParser()
    argument_parser.add_argument("--clean",help="",action='store_true')
    args = argument_parser.parse_args()
    
    delete_directory("dataset") # deletes existing directory if exists
    os.makedirs("dataset",exist_ok=True)
    df=pd.read_excel(r'files//export.xlsx')
    
    
    new_df=pd.DataFrame()
    dialogs_dataset=[]
    dialogs_dataset_seperated={key:[] for key in df["SubLanguageID"].unique()}
    file_path_language={key:[] for key in dialogs_dataset_seperated.keys()}
    for language in dialogs_dataset_seperated:
        os.makedirs(f"dataset/{language}",exist_ok=True)
    errors=[]
    fixed_files=set(os.listdir("time_fixed"))
    unique_ids=set()
    with tqdm(total=len(df)) as progress_bar:
        
        for index, row in df.iterrows():
            #id
            try:
                id=row["IDSubtitleFile"]
                #print(id)
                id=str(id)
                
                    #metadata
                language=row["SubLanguageID"]
                year=row["MovieYear"]
                idSubtitleFile=row["IDSubtitleFile"]
                subActualCD=row["SubActualCD"]
                subSumCD=row["SubSumCD"]
                subFormat=row["SubFormat"]
                movieName=row["MovieName"]
                movieimdbID=row["MovieImdbID"]
                userRank=row["UserRank"]
                subDownloadscnt=row["SubDownloadsCnt"]
                seriesimdbParent=row["SeriesIMDBParent"]
                seriesSeason=row["SeriesSeason"]
                seriesEpisode=row["SeriesEpisode"]
                subhearingImpaired=row["SubHearingImpaired"]
                encoding=row["Encoding"]
                
                file_path=reverse_id(id)
                file_path=find_file_path(file_path,id)
                if id[0]!='1':
                    id='1'+id 
                if id+".srt" in fixed_files:
                    file_path="time_fixed//"+id+".srt"
                    subtitles=parser.parse(file_path)
                
                subtitles = parser.parse(file_path)
                unique_ids.add(id)
                #dialogs = [dialog for dialog in subtitles]
                file_path_language[language].append(file_path)
                datas=[]
                
                if args.clean:
                    subtitles = brackets.clean(
                        formatting.clean(
                            urls_cleaner.clean(
                                subtitles
                            )
                        )
                    )
                
                
                
                
                
                for dialog in subtitles:
                    if dialog.text!="":
                        datas.append(dialog.text)
                        stats_language_wise[language]['dialogs']+=1
                        stats_language_wise[language]['total_characters']+=len(dialog.text)
                        stats_language_wise[language]['total_words']+=len(dialog.text.split(' '))
                        stats["dialogs"]+=1
                        stats['total_characters']+=len(dialog.text)
                        stats['total_words']+=len(dialog.text.split(' '))
                record={
                    
                        "ID":id,
                        "metadata":{
                            "MovieName":movieName,
                            "Year":year,
                            "IDSubtitleFile":idSubtitleFile,
                            "SubActualCD":subActualCD,
                            "SubSumCD":subSumCD,
                            "SubFormat":subFormat,
                            "MovieImdbID":movieimdbID,
                            "UserRank":userRank,
                            "SubDownloadsCnt":subDownloadscnt,
                            "SeriesIMDBParent":seriesimdbParent,
                            "SeriesSeason":seriesSeason,
                            "SeriesEpisode":seriesEpisode,
                            "SubHearingImpaired":subhearingImpaired,
                            "Encoding":encoding
                                
                        },
                        "dialogs":{
                            f"{language}":datas
                        }
                    }
                dialogs_dataset.append(
                    record
                )
                dialogs_dataset_seperated[language].append(
                    record
                )
                
            except Exception as e:
                errors.append(f"{e} at id: {id} FILE PATH: {file_path}")
            progress_bar.update(1)
            #if progress_bar.n==2:
            #    break
            
    try:
        pd.DataFrame(dialogs_dataset).to_json(path_or_buf="dataset/output.jsonl",lines=True,orient='records')
        for language in dialogs_dataset_seperated.keys():
            pd.DataFrame(dialogs_dataset_seperated[language]).to_json(path_or_buf=f"dataset/{language}/{language}.jsonl",lines=True,orient='records')
        
        for language in dialogs_dataset_seperated.keys():
            os.makedirs(f"dataset/{language}/subtitles",exist_ok=True)
        
        print(f"Copying files...")
        for language in dialogs_dataset_seperated.keys():
            dfolder=f"dataset//{language}//subtitles"
            print(f"Copying files {language}--->")
            for path in file_path_language[language]:
                source_path = path
                destination_path = dfolder+"//"+path.split("//")[-1]
                shutil.copy(source_path, destination_path)
            print("DONE")
    except Exception as e:
        print(f"Error: {e}")
    finally:
        for e in errors:
            print(e)
        stats['files']=len(df)
        for key in language_full_names.keys():
            stats_language_wise[key]['files']=(df['SubLanguageID'] == f'{key}').sum()
        print(stats)
        print(stats_language_wise)
        
    # rename everything 
       
    language_full_names = {
        "asm": "Assamese",
        "ben": "Bengali",
        "hin": "Hindi",
        "mar": "Marathi",
        "npl": "Nepali",
        "ori": "Odia",
        "snd": "Sindhi",
        "tam": "Tamil",
        "tel": "Telugu",
        "urd": "Urdu"
    }
    print("RENAMING FILES...")
    for folder in os.listdir("dataset"):
        if folder in language_full_names:
            new_name=language_full_names[folder]
            os.rename(f"dataset/{folder}/{folder}.jsonl",f"dataset/{folder}/{new_name}.jsonl")
            os.rename(f"dataset/{folder}",f"dataset/{new_name}")
            
    create_report(stats,stats_language_wise)
    print(f"Total unique File: {len(unique_ids)}")